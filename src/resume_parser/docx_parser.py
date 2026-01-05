"""
DOCX Parser для извлечения текста из Word документов
"""
import logging
from pathlib import Path
from typing import Optional, List
from docx import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DOCXParser:
    """Парсер для DOCX файлов"""

    def __init__(self):
        """Инициализация DOCX парсера"""
        self.supported_extensions = ['.docx', '.doc']

    def can_parse(self, file_path: str) -> bool:
        """
        Проверка, может ли парсер обработать файл

        Args:
            file_path: Путь к файлу

        Returns:
            True если файл можно обработать
        """
        ext = Path(file_path).suffix.lower()
        # .doc требует дополнительных библиотек, пока поддерживаем только .docx
        return ext == '.docx'

    def extract_text(self, file_path: str) -> Optional[str]:
        """
        Извлечение текста из DOCX файла

        Args:
            file_path: Путь к DOCX файлу

        Returns:
            Извлеченный текст или None при ошибке
        """
        if not self.can_parse(file_path):
            logger.error(f"Unsupported file type: {file_path}")
            return None

        try:
            path = Path(file_path)
            if not path.exists():
                logger.error(f"File not found: {file_path}")
                return None

            # Открываем документ
            doc = Document(path)

            # Извлекаем текст из параграфов
            text_parts = []

            # Параграфы
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Таблицы
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(table_text)

            # Объединяем весь текст
            full_text = "\n".join(text_parts)

            if not full_text.strip():
                logger.warning(f"No text extracted from {file_path}")
                return None

            logger.info(f"Successfully extracted {len(full_text)} characters from {path.name}")
            return full_text

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            return None

    def _extract_table_text(self, table) -> str:
        """
        Извлечение текста из таблицы

        Args:
            table: Объект таблицы python-docx

        Returns:
            Текст из таблицы
        """
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                table_text.append(" | ".join(row_text))

        return "\n".join(table_text)

    def extract_paragraphs(self, file_path: str) -> List[str]:
        """
        Извлечение параграфов как отдельного списка

        Args:
            file_path: Путь к DOCX файлу

        Returns:
            Список параграфов
        """
        paragraphs = []

        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())

        except Exception as e:
            logger.error(f"Error extracting paragraphs: {e}")

        return paragraphs

    def get_metadata(self, file_path: str) -> dict:
        """
        Получение метаданных DOCX

        Args:
            file_path: Путь к DOCX файлу

        Returns:
            Словарь с метаданными
        """
        metadata = {
            'num_paragraphs': 0,
            'num_tables': 0,
            'author': None,
            'title': None,
            'subject': None,
            'created': None,
            'modified': None
        }

        try:
            doc = Document(file_path)

            metadata['num_paragraphs'] = len(doc.paragraphs)
            metadata['num_tables'] = len(doc.tables)

            # Core properties
            if doc.core_properties:
                metadata.update({
                    'author': doc.core_properties.author,
                    'title': doc.core_properties.title,
                    'subject': doc.core_properties.subject,
                    'created': doc.core_properties.created,
                    'modified': doc.core_properties.modified
                })

        except Exception as e:
            logger.error(f"Error getting DOCX metadata: {e}")

        return metadata
